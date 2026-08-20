"""將 `assembly.assemble_document()` 的組裝結果，渲染成一份正式的企業
DOCX 檔案（Issue #9 What to build：封面／標題層級／頁首頁尾／表格／圖片
尺寸／圖說／頁碼）。

刻意獨立成一個模組（而不是塞進 views.py）：渲染邏輯本身不碰 HTTP，純粹
是「(assembled dict, project) -> bytes」的轉換函式，方便未來若要另外寫
渲染邏輯的單元測試（雖然 Issue #9 的驗收方式是走 HTTP 層 + 重新解析輸出，
但保持這個函式本身跟 Django view/request 物件無關，仍是比較乾淨的切法）。
"""

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from features.models import FeatureNode
from projects.models import PROJECT_VARIABLE_FIELDS, Project
from screenshots.models import ProjectScreenshot

# 標題層級上限——節點樹深度不限（Issue #1：FeatureNode 自我參照樹狀結構，
# 任意深度），但 Word 常見企業文件排版習慣不會無限往下開新的標題樣式，
# 深度超過此上限的節點一律沿用最深一級標題樣式，不再往下開新層級。
_MAX_HEADING_LEVEL = 4

_IMAGE_WIDTH = Inches(6)


def _node_depth(node_id: int, cache: dict) -> int:
    """計算 `node_id` 在 FeatureNode 樹中的深度（根節點＝0）。用 `cache`
    做記憶化，避免同一份文件裡重複節點的祖先鏈被重新查詢多次。"""
    if node_id in cache:
        return cache[node_id]

    node = FeatureNode.objects.filter(pk=node_id).only("id", "parent_id").first()
    if node is None or node.parent_id is None:
        depth = 0
    else:
        depth = 1 + _node_depth(node.parent_id, cache)

    cache[node_id] = depth
    return depth


def _add_page_number_field(paragraph):
    """在 `paragraph` 插入一個 Word「PAGE」動態欄位（頁碼），而不是寫死的
    數字文字——這樣即使實際分頁數量因內容多寡而變動，頁碼仍然正確（Word
    開啟後會自動計算，多數 DOCX 檢視器/PDF 轉檔工具也支援這個標準欄位）。
    """
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _setup_header_footer(document: Document, project: Project, document_type_label: str) -> None:
    section = document.sections[0]

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.text = (
        f"{project.customer_name} / {project.project_name} — {document_type_label}"
    )
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.add_run("第 ")
    _add_page_number_field(footer_paragraph)
    footer_paragraph.add_run(" 頁")


def _add_cover_page(document: Document, project: Project, document_type_label: str, generated_at) -> None:
    title = document.add_heading(document_type_label, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph(f"{project.customer_name} / {project.project_name}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)

    document.add_paragraph()

    # 表格 1：專案變數摘要（Issue #9 What to build：輸出需保留「表格」；
    # 這是最自然、跟 DOCX 內容本身相關的表格用途——把封面上該有的專案基本
    # 資料以表格列出，而不是塞成一長串散文）。
    table = document.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for field_name in PROJECT_VARIABLE_FIELDS:
        field = Project._meta.get_field(field_name)
        value = getattr(project, field_name)
        if field_name == "acceptance_date":
            display_value = value.isoformat() if value else ""
        else:
            display_value = value or ""
        row_cells = table.add_row().cells
        row_cells[0].text = str(field.verbose_name)
        row_cells[1].text = display_value

    document.add_paragraph()
    meta = document.add_paragraph(f"產生時間：{generated_at.strftime('%Y-%m-%d %H:%M:%S')}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_prose_paragraphs(document: Document, body: str) -> None:
    if not body:
        return
    # 內容欄位本身是一大段散文文字（`FeatureNodeContent.body` /
    # override 的 body），以空白行分段——讓後台編輯者用一般的段落分段習慣
    # 寫內容時，DOCX 輸出也能正確斷成多個段落，而不是整段擠成一坨。
    for paragraph_text in body.split("\n\n"):
        stripped = paragraph_text.strip()
        if stripped:
            document.add_paragraph(stripped)


def _add_screenshot(document: Document, screenshot_entry: dict) -> None:
    if not screenshot_entry.get("applicable"):
        return
    screenshot_ref = screenshot_entry.get("screenshot")
    if not screenshot_ref:
        return

    shot = ProjectScreenshot.objects.filter(pk=screenshot_ref["id"]).first()
    if shot is None or not shot.image:
        return

    shot.image.open("rb")
    try:
        image_bytes = shot.image.read()
    finally:
        shot.image.close()

    document.add_picture(io.BytesIO(image_bytes), width=_IMAGE_WIDTH)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption_text = screenshot_entry.get("caption") or shot.original_filename
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(f"圖：{caption_text}")
    caption_run.italic = True
    caption_run.font.size = Pt(10)


def _add_section(document: Document, section: dict, depth_cache: dict) -> None:
    if section["type"] == "node":
        depth = _node_depth(section["node"], depth_cache)
        level = min(depth + 1, _MAX_HEADING_LEVEL)
    else:
        level = 1

    document.add_heading(section["title"], level=level)
    _add_prose_paragraphs(document, section.get("body", ""))

    for screenshot_entry in section.get("screenshots", []):
        _add_screenshot(document, screenshot_entry)


def build_docx(assembled: dict, project: Project, generated_at) -> bytes:
    """把 `assemble_document()` 的回傳值（`assembled`）渲染成一份 DOCX
    檔案的完整 bytes。只渲染 `included=True` 的章節——已被本文件類型預覽
    層級排除的章節，不應該出現在最終產出的檔案裡（跟
    `validate_generation_readiness()` 判斷「這份文件真正會擋下產生的缺漏
    項目」用同一個 `included` 旗標，維持一致）。
    """
    document = Document()
    document_type_label = assembled["document_type_label"]

    _setup_header_footer(document, project, document_type_label)
    _add_cover_page(document, project, document_type_label, generated_at)
    document.add_page_break()

    depth_cache: dict = {}
    for section in assembled["sections"]:
        if not section["included"]:
            continue
        _add_section(document, section, depth_cache)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
