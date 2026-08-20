"""HTTP-layer tests for T8: DOCX 文件產生與歷史紀錄 (Issue #9).

Per the MVP spec's Testing Decisions, the HTTP API layer is the only test
seam. For the binary DOCX output specifically, the HTTP-layer-appropriate
test is: call the generate endpoint, take the returned file bytes, and
parse them back with `python-docx` to assert on observable structure
(heading text present, excluded content absent, an actual image embedded)
— this exercises the real output of the HTTP endpoint, not an internal
implementation detail.
"""

import io
import shutil
import struct
import tempfile
import zlib

from django.contrib.auth.models import User
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import Client, TestCase, override_settings
from docx import Document as DocxDocument

from features.models import FeatureNode, FeatureNodeContent
from projects.models import Project
from screenshots.models import ProjectScreenshot, ScreenshotRequirement
from selections.models import ProjectFeatureExclusion, ProjectFeatureSelection

from ..models import GeneratedDocument

_TEST_MEDIA_ROOT = tempfile.mkdtemp(prefix="gen_doc_documents_test_")


def tearDownModule():
    shutil.rmtree(_TEST_MEDIA_ROOT, ignore_errors=True)


# --- 手刻一張「真的合法」的最小 PNG（純標準函式庫，不依賴 Pillow——本專
# 案刻意不把 Pillow 列為依賴，見 screenshots/models.py 的既有判斷）。
# python-docx 的 `add_picture()` 無論有沒有指定 width/height，都一定要能
# 解析圖片檔頭來判斷格式與原生尺寸，所以測試用圖不能只是「PNG signature +
# 隨便填充 bytes」（screenshots app 既有測試的 `_valid_png()` 就是這種、
# 只夠通過 view 層的 magic-bytes 簽章驗證，不足以被 python-docx 解析），
# 必須是結構完整（IHDR/IDAT/IEND 三個 chunk、正確 CRC）的最小合法 PNG。
def _png_chunk(chunk_type: bytes, data: bytes) -> bytes:
    return (
        struct.pack(">I", len(data))
        + chunk_type
        + data
        + struct.pack(">I", zlib.crc32(chunk_type + data) & 0xFFFFFFFF)
    )


def _valid_png_bytes(width=20, height=20) -> bytes:
    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = _png_chunk(
        b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    )
    raw_scanlines = b"".join(
        b"\x00" + bytes([255, 0, 0]) * width for _ in range(height)
    )
    idat = _png_chunk(b"IDAT", zlib.compress(raw_scanlines))
    iend = _png_chunk(b"IEND", b"")
    return signature + ihdr + idat + iend


def _valid_png_upload(name="shot.png"):
    return SimpleUploadedFile(name, _valid_png_bytes(), content_type="image/png")


def _generate_url(project_id, document_type):
    return f"/api/projects/{project_id}/documents/{document_type}/generate/"


def _history_url(project_id, document_type):
    return f"/api/projects/{project_id}/documents/{document_type}/history/"


def _status_url():
    return "/api/projects/documents-status/"


def _all_paragraph_text(docx_document: DocxDocument) -> str:
    return "\n".join(p.text for p in docx_document.paragraphs)


def _embedded_image_blob_sizes(docx_document: DocxDocument) -> list:
    sizes = []
    for shape in docx_document.inline_shapes:
        rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        image_part = docx_document.part.related_parts[rId]
        sizes.append(len(image_part.blob))
    return sizes


class AuthGateTests(TestCase):
    def setUp(self):
        owner = User.objects.create_user(username="owner1", password="pw12345")
        self.project = Project.objects.create(
            owner=owner, customer_name="X", project_name="Y"
        )

    def test_generate_requires_auth(self):
        response = self.client.post(_generate_url(self.project.id, "system_test"))
        self.assertEqual(response.status_code, 401)

    def test_history_requires_auth(self):
        response = self.client.get(_history_url(self.project.id, "system_test"))
        self.assertEqual(response.status_code, 401)

    def test_status_requires_auth(self):
        response = self.client.get(_status_url())
        self.assertEqual(response.status_code, 401)


class InvalidDocumentTypeTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="pm_invalid", password="pw12345")
        self.client.login(username="pm_invalid", password="pw12345")
        self.project = Project.objects.create(
            owner=self.user, customer_name="客戶", project_name="專案"
        )

    def test_unknown_document_type_generate_is_404(self):
        response = self.client.post(_generate_url(self.project.id, "not_a_real_type"))
        self.assertEqual(response.status_code, 404)

    def test_unknown_document_type_history_is_404(self):
        response = self.client.get(_history_url(self.project.id, "not_a_real_type"))
        self.assertEqual(response.status_code, 404)

    def test_training_deck_generate_is_blocked_with_400(self):
        # training_deck (PPTX) 是 T9 範圍，T8 明確擋下、回傳 400（不是
        # 404——這是「已知但尚未支援」，跟「根本不存在的代碼」語意不同）。
        response = self.client.post(_generate_url(self.project.id, "training_deck"))
        self.assertEqual(response.status_code, 400)
        self.assertEqual(GeneratedDocument.objects.count(), 0)

    def test_training_deck_history_returns_empty_list_not_error(self):
        # 歷史查詢端點刻意對四種文件類型都通用（見 views.py 設計判斷）。
        response = self.client.get(_history_url(self.project.id, "training_deck"))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["results"], [])


@override_settings(MEDIA_ROOT=_TEST_MEDIA_ROOT)
class GenerationBlockedByReadinessTests(TestCase):
    """Issue #9 What to build：產生前必須呼叫
    `validate_generation_readiness()`，未 ready 時擋下產生，回傳明確列出
    缺漏項目的錯誤。"""

    def setUp(self):
        self.user = User.objects.create_user(username="pm_block", password="pw12345")
        self.client.login(username="pm_block", password="pw12345")
        self.project = Project.objects.create(
            owner=self.user, customer_name="客戶", project_name="專案"
        )
        self.node = FeatureNode.objects.create(name="帳號管理")
        FeatureNodeContent.objects.create(
            node=self.node,
            document_type="system_test",
            body="系統版本：{{project_variables.system_version}}",
        )
        self.requirement = ScreenshotRequirement.objects.create(
            node=self.node, name="帳號列表"
        )
        ProjectFeatureSelection.objects.create(project=self.project, node=self.node)

    def test_missing_screenshot_and_variable_blocks_generation(self):
        response = self.client.post(_generate_url(self.project.id, "system_test"))
        self.assertEqual(response.status_code, 400)
        body = response.json()
        self.assertEqual(len(body["missing_screenshots"]), 1)
        self.assertEqual(body["missing_screenshots"][0]["requirement_id"], self.requirement.id)
        self.assertIn("system_version", body["missing_variables"])
        self.assertEqual(GeneratedDocument.objects.count(), 0)


@override_settings(MEDIA_ROOT=_TEST_MEDIA_ROOT)
class DocxContentTests(TestCase):
    """對三種 DOCX 文件類型各自驗證：章節內容確實出現、未勾選/被排除節點
    的內容不出現、截圖確實被嵌入（非空/非佔位資料）——Issue #9 acceptance
    criterion 1 與「Testing the actual DOCX output」的要求。"""

    def _build_project(self, document_type):
        user = User.objects.create_user(
            username=f"pm_{document_type}", password="pw12345"
        )
        self.client.login(username=user.username, password="pw12345")
        project = Project.objects.create(
            owner=user,
            customer_name="測試客戶",
            project_name="測試專案",
            system_url="https://portal.example.com",
        )

        included_root = FeatureNode.objects.create(name="帳號管理")
        included_child = FeatureNode.objects.create(
            name="編輯帳號", parent=included_root
        )
        excluded_child = FeatureNode.objects.create(
            name="刪除帳號確認", parent=included_root
        )
        unchecked_root = FeatureNode.objects.create(name="報表")

        for node, body in (
            (included_root, "帳號管理章節內容，系統網址：{{project_variables.system_url}}"),
            (included_child, "編輯帳號章節內容"),
            (excluded_child, "刪除帳號確認章節內容－不應出現"),
            (unchecked_root, "報表章節內容－不應出現"),
        ):
            FeatureNodeContent.objects.create(
                node=node, document_type=document_type, body=body
            )

        ProjectFeatureSelection.objects.create(project=project, node=included_root)
        ProjectFeatureExclusion.objects.create(project=project, node=excluded_child)

        requirement = ScreenshotRequirement.objects.create(
            node=included_root, name="帳號列表"
        )
        ProjectScreenshot.objects.create(
            project=project,
            requirement=requirement,
            image=_valid_png_upload(),
            original_filename="shot.png",
            document_types=[document_type],
            is_custom=False,
            caption="帳號列表畫面截圖",
        )

        return project

    def _assert_docx_generation(self, document_type):
        project = self._build_project(document_type)

        response = self.client.post(_generate_url(project.id, document_type))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("attachment", response["Content-Disposition"])
        self.assertTrue(len(response.content) > 0)

        record_id = int(response["X-Generated-Document-Id"])
        record = GeneratedDocument.objects.get(pk=record_id)
        self.assertEqual(record.project_id, project.id)
        self.assertEqual(record.document_type, document_type)
        self.assertEqual(record.output_format, "docx")

        docx_document = DocxDocument(io.BytesIO(response.content))
        all_text = _all_paragraph_text(docx_document)

        # 已勾選章節的內容出現，且變數已被替換。
        self.assertIn("帳號管理章節內容", all_text)
        self.assertIn("https://portal.example.com", all_text)
        self.assertNotIn("{{project_variables", all_text)
        self.assertIn("編輯帳號章節內容", all_text)

        # 未勾選節點、被排除的子孫節點內容都不出現。
        self.assertNotIn("刪除帳號確認章節內容", all_text)
        self.assertNotIn("報表章節內容", all_text)

        # 封面：客戶/專案名稱、標題層級（heading）存在。
        self.assertIn("測試客戶", all_text)
        self.assertIn("測試專案", all_text)
        heading_styles = {p.style.name for p in docx_document.paragraphs if p.text.strip()}
        self.assertTrue(any(name.startswith("Heading") for name in heading_styles))

        # 表格（封面的專案變數摘要）。
        self.assertTrue(len(docx_document.tables) >= 1)

        # 截圖確實被嵌入：至少一張圖，且是真的圖片資料（非空/非佔位）。
        blob_sizes = _embedded_image_blob_sizes(docx_document)
        self.assertEqual(len(blob_sizes), 1)
        self.assertGreater(blob_sizes[0], 50)
        self.assertIn("帳號列表畫面截圖", all_text)

        # 頁首含客戶/專案資訊；頁尾含頁碼相關文字。
        header_text = docx_document.sections[0].header.paragraphs[0].text
        self.assertIn("測試客戶", header_text)
        footer_text = docx_document.sections[0].footer.paragraphs[0].text
        self.assertIn("頁", footer_text)

        return project, record

    def test_system_design_document_generation(self):
        self._assert_docx_generation("system_design")

    def test_system_test_document_generation(self):
        self._assert_docx_generation("system_test")

    def test_system_install_document_generation(self):
        self._assert_docx_generation("system_install")


@override_settings(MEDIA_ROOT=_TEST_MEDIA_ROOT)
class RegenerationHistoryTests(TestCase):
    """Issue #9 acceptance criteria 2、3：重新產生新增新的一筆歷史紀錄
    （不覆蓋前一筆），新紀錄反映新值、先前紀錄內容不變；查詢歷史可看到全
    部紀錄與各自產生時間，預設只呈現最新一筆。"""

    def setUp(self):
        self.user = User.objects.create_user(username="pm_hist", password="pw12345")
        self.client.login(username="pm_hist", password="pw12345")
        self.project = Project.objects.create(
            owner=self.user, customer_name="舊客戶名稱", project_name="專案"
        )
        self.node = FeatureNode.objects.create(name="登入")
        FeatureNodeContent.objects.create(
            node=self.node, document_type="system_test", body="登入章節內容"
        )
        ProjectFeatureSelection.objects.create(project=self.project, node=self.node)

    def _generate(self):
        response = self.client.post(_generate_url(self.project.id, "system_test"))
        self.assertEqual(response.status_code, 200)
        return response

    def test_regenerating_after_customer_name_change_creates_new_row_old_row_unchanged(self):
        first_response = self._generate()
        first_id = int(first_response["X-Generated-Document-Id"])

        self.project.customer_name = "新客戶名稱"
        self.project.save(update_fields=["customer_name"])

        second_response = self._generate()
        second_id = int(second_response["X-Generated-Document-Id"])

        self.assertNotEqual(first_id, second_id)
        self.assertEqual(GeneratedDocument.objects.filter(project=self.project).count(), 2)

        first_record = GeneratedDocument.objects.get(pk=first_id)
        second_record = GeneratedDocument.objects.get(pk=second_id)

        self.assertEqual(first_record.content_snapshot["sections"][0]["body"], "登入章節內容")

        # 先前紀錄的快照 customer_name 相關資訊不受後續修改影響——直接重新
        # 解析先前紀錄實際存的 DOCX 檔案內容驗證，而不是只看 DB row 數。
        old_docx = DocxDocument(io.BytesIO(first_response.content))
        self.assertIn("舊客戶名稱", _all_paragraph_text(old_docx))
        self.assertNotIn("新客戶名稱", _all_paragraph_text(old_docx))

        new_docx = DocxDocument(io.BytesIO(second_response.content))
        self.assertIn("新客戶名稱", _all_paragraph_text(new_docx))
        self.assertNotIn("舊客戶名稱", _all_paragraph_text(new_docx))

        # 重新讀取先前紀錄實際落地儲存的檔案本身，確認也未被覆蓋。
        first_record.file.open("rb")
        try:
            stored_old_bytes = first_record.file.read()
        finally:
            first_record.file.close()
        stored_old_docx = DocxDocument(io.BytesIO(stored_old_bytes))
        self.assertIn("舊客戶名稱", _all_paragraph_text(stored_old_docx))
        self.assertNotIn("新客戶名稱", _all_paragraph_text(stored_old_docx))

    def test_history_endpoint_returns_all_three_generations_latest_default(self):
        first = self._generate()
        first_id = int(first["X-Generated-Document-Id"])
        second = self._generate()
        second_id = int(second["X-Generated-Document-Id"])
        third = self._generate()
        third_id = int(third["X-Generated-Document-Id"])

        response = self.client.get(_history_url(self.project.id, "system_test"))
        self.assertEqual(response.status_code, 200)
        results = response.json()["results"]

        self.assertEqual(len(results), 3)
        self.assertEqual({r["id"] for r in results}, {first_id, second_id, third_id})

        # 新到舊排序，且只有最新一筆標示 is_latest。
        self.assertEqual(results[0]["id"], third_id)
        self.assertTrue(results[0]["is_latest"])
        self.assertFalse(results[1]["is_latest"])
        self.assertFalse(results[2]["is_latest"])

        # 各自的產生時間都存在。
        timestamps = [r["generated_at"] for r in results]
        for ts in timestamps:
            self.assertTrue(ts)

        # 每一筆都能取得完整快照。
        for r in results:
            self.assertIn("sections", r["content_snapshot"])


@override_settings(MEDIA_ROOT=_TEST_MEDIA_ROOT)
class ProjectListStatusTests(TestCase):
    """Issue #9 acceptance criterion 4：專案列表可看到各專案這三類 DOCX
    文件目前是否已產生。"""

    def setUp(self):
        self.user = User.objects.create_user(username="pm_status", password="pw12345")
        self.client.login(username="pm_status", password="pw12345")
        self.project_a = Project.objects.create(
            owner=self.user, customer_name="A客戶", project_name="A專案"
        )
        self.project_b = Project.objects.create(
            owner=self.user, customer_name="B客戶", project_name="B專案"
        )
        self.node = FeatureNode.objects.create(name="登入")
        FeatureNodeContent.objects.create(
            node=self.node, document_type="system_test", body="內容"
        )
        ProjectFeatureSelection.objects.create(project=self.project_a, node=self.node)
        ProjectFeatureSelection.objects.create(project=self.project_b, node=self.node)

    def test_status_reflects_per_document_type_generation_state(self):
        # 尚未產生任何文件前，全部為 False。
        response = self.client.get(_status_url())
        results_by_project = {r["project"]: r for r in response.json()["results"]}
        self.assertFalse(results_by_project[self.project_a.id]["system_test"])
        self.assertFalse(results_by_project[self.project_a.id]["system_design"])
        self.assertFalse(results_by_project[self.project_a.id]["system_install"])

        # 只替 project_a 產生 system_test。
        gen_response = self.client.post(_generate_url(self.project_a.id, "system_test"))
        self.assertEqual(gen_response.status_code, 200)

        response = self.client.get(_status_url())
        results_by_project = {r["project"]: r for r in response.json()["results"]}

        self.assertTrue(results_by_project[self.project_a.id]["system_test"])
        self.assertFalse(results_by_project[self.project_a.id]["system_design"])
        self.assertFalse(results_by_project[self.project_a.id]["system_install"])

        # project_b 完全不受影響。
        self.assertFalse(results_by_project[self.project_b.id]["system_test"])
        self.assertFalse(results_by_project[self.project_b.id]["system_design"])
        self.assertFalse(results_by_project[self.project_b.id]["system_install"])


@override_settings(MEDIA_ROOT=_TEST_MEDIA_ROOT)
class GenerateCsrfProtectionTests(TestCase):
    """產生端點是會改變狀態的寫入端點，必須受 CSRF 保護（絕不可
    `@csrf_exempt`——同 T7 對 adjustments 端點的既有慣例）。"""

    def setUp(self):
        self.user = User.objects.create_user(username="pm_csrf", password="pw12345")
        self.project = Project.objects.create(
            owner=self.user, customer_name="客戶", project_name="專案"
        )
        self.node = FeatureNode.objects.create(name="登入")
        FeatureNodeContent.objects.create(
            node=self.node, document_type="system_test", body="內容"
        )
        ProjectFeatureSelection.objects.create(project=self.project, node=self.node)
        self.client = Client(enforce_csrf_checks=True)
        self.client.login(username="pm_csrf", password="pw12345")

    def _csrf_token(self):
        return self.client.get("/api/auth/csrf/").json()["csrfToken"]

    def test_generate_without_csrf_token_is_rejected(self):
        response = self.client.post(_generate_url(self.project.id, "system_test"))
        self.assertEqual(response.status_code, 403)
        self.assertEqual(GeneratedDocument.objects.count(), 0)

    def test_generate_with_primed_csrf_token_succeeds(self):
        token = self._csrf_token()
        response = self.client.post(
            _generate_url(self.project.id, "system_test"),
            HTTP_X_CSRFTOKEN=token,
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(GeneratedDocument.objects.count(), 1)
